import json
from docx import Document

# Load JSON file
with open("CMS_QPs.json", "r") as f:
    qp_data = json.load(f)

# Open Word template
doc = Document("finaltemplete.docx")

# Build replacement dictionary
replacements = {}

for key, value in qp_data.items():
    # Question text
    replacements[f"{{{{{key}}}}}"] = value["question"]
    # CO text
    replacements[f"{{{{CO_{key}}}}}"] = value["co"]

# Replace in normal paragraphs
for para in doc.paragraphs:
    for k, v in replacements.items():
        if k in para.text:
            para.text = para.text.replace(k, v)

# Replace in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for k, v in replacements.items():
                    if k in para.text:
                        para.text = para.text.replace(k, v)

# Save output
doc.save("CMS_output_paper.docx")

print("Question paper generated with CO and questions")
